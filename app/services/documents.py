from __future__ import annotations

import re
from pathlib import Path

from docx import Document


def safe_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_") or "application"


def export_docx(company: str, role: str, kind: str, text: str, directory: Path) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"{safe_filename(company)}_{safe_filename(role)}_{safe_filename(kind)}.docx"
    document = Document()
    document.add_heading("Shreyas Dattashivarama", level=0)
    document.add_heading(kind.replace("_", " "), level=1)
    for paragraph in text.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(path)
    return path

